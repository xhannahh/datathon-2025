import pymupdf as fitz  # PyMuPDF
import docx
from typing import Dict, List, Tuple
from PIL import Image
import io
import base64
import uuid, os, cv2, fitz, numpy as np
import pytesseract
from pytesseract import Output

default_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
pytesseract.pytesseract.tesseract_cmd = os.getenv("TESSERACT_CMD", default_path)


def extract_from_pdf(path: str) -> Tuple[Dict[int, str], int, float, List[Dict]]:
   """Extract text and images from PDF.
  
   Returns:
       Tuple of (pages_text, image_count, images_data)
       where images_data is a list of dicts with 'page', 'index', and 'data' (base64)
   """
   doc = fitz.open(path)
   pages = {}
   images_data = []
   image_count = 0
  
   for i, page in enumerate(doc, start=1):
       pages[i] = page.get_text("text") or ""
       image_list = page.get_images(full=True)
      
       for img_index, img_info in enumerate(image_list):
           try:
               xref = img_info[0]
               base_image = doc.extract_image(xref)
               image_bytes = base_image["image"]
               image_ext = base_image["ext"]
              
               # Convert to base64 for storage and API calls
               image_b64 = base64.b64encode(image_bytes).decode('utf-8')
              
               images_data.append({
                   "page": i,
                   "index": img_index,
                   "data": image_b64,
                   "ext": image_ext,
                   "size": len(image_bytes)
               })
               image_count += 1
           except Exception as e:
               # Skip images that can't be extracted
               print(f"Failed to extract image {img_index} from page {i}: {e}")
               continue


   legibility_score = analyze_pdf_legibility(path)
   legibility_report = (
   sum(s["combined_legibility"] for s in legibility_score) / len(legibility_score)
   if legibility_score else 0.0
   )
              
   return pages, image_count, legibility_report, images_data


def extract_from_docx(path: str) -> Tuple[Dict[int, str], int, float, List[Dict]]:
   """Extract text, images, and legibility from DOCX using pytesseract for OCR."""
   document = docx.Document(path)
   pages = _split_docx_into_pages(document)
   full_text = "\n".join(pages.values()).strip()

   images_data, legibility_scores = _extract_docx_images(document)
   image_count = len(images_data)

   legibility_report = (
       round(sum(legibility_scores) / len(legibility_scores), 3)
       if legibility_scores else
       (1.0 if len(full_text) > 50 else 0.0)
   )

   return pages, image_count, legibility_report, images_data




def pdf_to_images(pdf_path: str, dpi: int = 150):
   """Convert all pages of a PDF into RGB images."""
   doc = fitz.open(pdf_path)
   images = []
   for page_num in range(len(doc)):
       page = doc.load_page(page_num)
       pix = page.get_pixmap(matrix=fitz.Matrix(dpi / 72, dpi / 72))
       img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, 3)
       img_rgb = cv2.cvtColor(img, cv2.COLOR_BGR2RGB)
       images.append(img_rgb)
   doc.close()
   return images




def sharpness_score(img_rgb: np.ndarray) -> float:
   """Compute image sharpness using variance of Laplacian."""
   gray = cv2.cvtColor(img_rgb, cv2.COLOR_RGB2GRAY)
   return float(cv2.Laplacian(gray, cv2.CV_64F).var())




def ocr_confidence_score(img_rgb: np.ndarray) -> float:
   """Compute average OCR confidence as proxy for text legibility."""
   data = pytesseract.image_to_data(img_rgb, output_type=Output.DICT)
   confs = []
   for c in data['conf']:
       try:
           val = int(c)
           if val > 0:
               confs.append(val)
       except (ValueError, TypeError):
           continue
   if not confs:
       return 0.0
   return sum(confs) / len(confs)




def combined_legibility(img_rgb: np.ndarray) -> float:
   """Blend image sharpness and OCR confidence into a single legibility score (0–1)."""
   sharp = sharpness_score(img_rgb)
   ocr_conf = ocr_confidence_score(img_rgb)
   # Normalize and weight
   sharp_norm = min(sharp / 1000, 1.0)
   ocr_norm = ocr_conf / 100.0
   return round(0.5 * sharp_norm + 0.5 * ocr_norm, 3)




def analyze_pdf_legibility(pdf_path: str):
   """Compute per-page legibility for a PDF."""
   pages = pdf_to_images(pdf_path)
   results = []
   for i, img in enumerate(pages, 1):
       leg_score = combined_legibility(img)
       results.append({
           "page": i,
           "sharpness": sharpness_score(img),
           "ocr_confidence": ocr_confidence_score(img),
           "combined_legibility": leg_score
       })
   return results


def extract_generic(path: str) -> Tuple[Dict[int, str], int, float, List[Dict]]:
   """Extract text and images from document.
  
   Returns:
       Tuple of (pages_text, image_count, images_data)
   """
   if path.lower().endswith(".pdf"):
       return extract_from_pdf(path)
   if path.lower().endswith(".docx"):
       return extract_from_docx(path)
   # fallback: treat as text
   with open(path, "r", errors="ignore") as f:
       return {1: f.read()}, 0, 0.0, []


def _split_docx_into_pages(document: docx.Document) -> Dict[int, str]:
   """Approximate DOCX pagination using explicit breaks and length-based fallback."""
   pages: Dict[int, str] = {}
   page_num = 1
   current_lines: List[str] = []

   for para in document.paragraphs:
       text = para.text.strip()
       has_break = _has_page_break(para)

       if text:
           current_lines.append(text)

       if has_break:
           page_text = "\n".join(current_lines).strip()
           if page_text:
               pages[page_num] = page_text
               page_num += 1
           current_lines = []

   if current_lines:
       page_text = "\n".join(current_lines).strip()
       if page_text:
           pages[page_num] = page_text

   if not pages:
       all_text = "\n".join(p.text for p in document.paragraphs).strip()
       pages = {1: all_text or ""}
   elif len(pages) == 1:
       all_text = pages[1]
       chars_per_page = 2200
       if len(all_text) > chars_per_page:
           pages = {}
           page_num = 1
           for start in range(0, len(all_text), chars_per_page):
               chunk = all_text[start:start + chars_per_page].strip()
               if chunk:
                   pages[page_num] = chunk
                   page_num += 1
           if not pages:
               pages = {1: all_text}

   return pages


def _has_page_break(para) -> bool:
   try:
       br = para._element.xpath('.//w:br[@w:type="page"]')
       sect = para._element.xpath('.//w:pPr/w:sectPr')
       return bool(br or sect)
   except Exception:
       return False


def _extract_docx_images(document: docx.Document) -> Tuple[List[Dict], List[float]]:
   """Extract embedded images from DOCX and compute legibility heuristics."""
   images_data: List[Dict] = []
   legibility_scores: List[float] = []

   for rel in document.part.rels.values():
       if getattr(rel, "is_external", False):
           continue
       try:
           target = rel.target_part
       except ValueError:
           continue
       if target is None:
           continue
       content_type = getattr(target, "content_type", "") or ""
       if not content_type.startswith("image/"):
           continue
       try:
           image_bytes = target.blob
           ext = content_type.split("/")[-1] or "png"
           image_b64 = base64.b64encode(image_bytes).decode("utf-8")
           index = len(images_data)
           images_data.append({
               "page": 1,  # python-docx does not expose precise pagination
               "index": index,
               "data": image_b64,
               "ext": ext,
               "size": len(image_bytes)
           })

           np_arr = np.frombuffer(image_bytes, np.uint8)
           img_rgb = cv2.imdecode(np_arr, cv2.IMREAD_COLOR)
           if img_rgb is not None:
               ocr_data = pytesseract.image_to_data(img_rgb, output_type=Output.DICT)
               confs: List[int] = []
               for c in ocr_data.get("conf", []):
                   try:
                       val = int(float(c))
                   except (ValueError, TypeError):
                       continue
                   if val > 0:
                       confs.append(val)
               ocr_conf = sum(confs) / len(confs) if confs else 0.0
               sharp = float(cv2.Laplacian(cv2.cvtColor(img_rgb, cv2.COLOR_BGR2GRAY), cv2.CV_64F).var())
               sharp_norm = min(sharp / 1000, 1.0)
               ocr_norm = ocr_conf / 100.0
               legibility_scores.append(0.5 * sharp_norm + 0.5 * ocr_norm)
       except Exception as exc:
           print(f"Failed to extract DOCX image: {exc}")
           continue

   return images_data, legibility_scores
